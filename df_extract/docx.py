from docx import Document

from df_extract.base import BaseExtract
from df_extract.extract_util import cleanup
from df_extract.utils import iter_to_aiter


class ExtractDocx(BaseExtract):
    """
    Class for Docx or Doc Extraction
    """

    async def extract_as_text(self, doc: Document):
        """
        Method to extract docx or doc content as text

        :param doc: Valid `docx.Document` object
        """
        await self.remove_existing_output()
        text = ''
        async for para in iter_to_aiter(doc.paragraphs):
            cleaned_up_data = await cleanup(para.text)
            text += cleaned_up_data + '\n\n\n\n'

        await self._write_text_output(text=text)

    async def extract_as_json(self, doc: Document):
        """
        Method to extract docx or doc content as json

        :param doc: Valid `docx.Document` object
        """
        await self.remove_existing_json_output()
        data = []
        _para_count = 1
        async for para in iter_to_aiter(doc.paragraphs):
            para_content = {
                'number': _para_count,
                'content': await cleanup(para.text),
                'name': self.name
            }
            data.append(para_content)
            _para_count += 1

        await self._write_json_output(data=data)

    async def extract(self, as_json: bool = False) -> None:
        """
        Method to extract csv content

        :param as_json: Extracted content will be stored as json. Default `False`
        """
        print(f'Extracting => {self.file_path}')
        doc = Document(self.file_path)
        if not as_json:
            await self.extract_as_text(doc)
        else:
            await self.extract_as_json(doc)
        print(f'Extracted => {self.file_path}')
